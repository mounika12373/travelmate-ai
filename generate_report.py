import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Ensure output directory exists
OUTPUT_PATH = "TravelMate_AI_Project_Report.docx"

def add_page_number(run):
    """Inserts a dynamic page number field into a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_custom_heading(doc, text, level, space_before=12, space_after=6):
    """Adds a styled heading with custom font, size, and colors."""
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    # Heading sizes and colors
    if level == 1:
        font_size = 18
        color = RGBColor(27, 54, 93)     # Deep Navy
        bold = True
    elif level == 2:
        font_size = 14
        color = RGBColor(112, 128, 144)  # Slate Gray
        bold = True
    else:
        font_size = 12
        color = RGBColor(51, 51, 51)     # Charcoal
        bold = True
        
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold
    return h

def add_custom_paragraph(doc, text, space_after=6, line_spacing=1.15, bold=False, italic=False):
    """Adds a standard body paragraph with custom spacing and fonts."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 51, 51)  # Charcoal
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_point(doc, bold_prefix, text, num_level=None):
    """Adds a list item (bullet or numbered) with custom formatting."""
    style_name = 'List Number' if num_level else 'List Bullet'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(51, 51, 51)
        r1.bold = True
        
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(51, 51, 51)
    return p

def create_styled_table(doc, rows, cols):
    """Creates a centered table with clean borders."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set default borders to thin light gray
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        r'<w:tblBorders {} >'
        r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        r'  <w:insideV w:val="none"/>'
        r'</w:tblBorders>'.format(nsdecls('w'))
    )
    tblPr.append(borders)
    return table

def style_table_header(row, titles):
    """Styles the header row of a table (Navy background, white bold text)."""
    # Set height
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))
    trPr.append(parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w'))))
    
    for idx, text in enumerate(titles):
        cell = row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        
        # Shading
        shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        # Vertical alignment center
        cell.vertical_alignment = 1

def style_table_row(row, values, is_even=False):
    """Styles a regular row in a table with zebra striping on even rows."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))
    
    for idx, text in enumerate(values):
        cell = row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 51, 51)
        
        if is_even:
            shading = parse_xml(r'<w:shd {} w:fill="F4F6F9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading)
        
        cell.vertical_alignment = 1

def add_figure(doc, image_path, caption, width_in_inches=5.5):
    """Inserts a centered image with a styled caption below it."""
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(4)
    
    if os.path.exists(image_path):
        try:
            run_img = p_img.add_run()
            run_img.add_picture(image_path, width=Inches(width_in_inches))
        except Exception as e:
            run_err = p_img.add_run(f"\n[Error loading image {image_path}: {e}]")
            run_err.font.color.rgb = RGBColor(180, 50, 50)
            run_err.bold = True
    else:
        run_miss = p_img.add_run(f"\n[Screenshot Not Found: {image_path}]")
        run_miss.font.color.rgb = RGBColor(180, 50, 50)
        run_miss.bold = True
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.keep_with_next = False
    
    run_cap = p_cap.add_run(caption)
    run_cap.font.name = 'Calibri'
    run_cap.font.size = Pt(9.5)
    run_cap.font.color.rgb = RGBColor(112, 128, 144)
    run_cap.italic = True

def main():
    doc = Document()
    
    # Configure 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default paragraph format
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    # ==========================================
    # 1. COVER PAGE
    # ==========================================
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)
    
    p_quote_title = doc.add_paragraph()
    p_quote_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_quote_title = p_quote_title.add_run('“TRAVELMATE AI: SMART TRAVEL COMPANION AND GUIDANCE PLATFORM”')
    r_quote_title.font.name = 'Arial'
    r_quote_title.font.size = Pt(22)
    r_quote_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    r_quote_title.bold = True
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(18)
    p_subtitle.paragraph_format.space_after = Pt(18)
    r_subtitle = p_subtitle.add_run(
        'A PROJECT REPORT\n'
        'Submitted in partial fulfillment of the requirements for the award of the degree of\n'
    )
    r_subtitle.font.size = Pt(12)
    r_subtitle.italic = True
    
    r_degree = p_subtitle.add_run('BACHELOR OF TECHNOLOGY\nIN\nARTIFICIAL INTELLIGENCE AND MACHINE LEARNING')
    r_degree.font.size = Pt(13)
    r_degree.bold = True
    r_degree.font.color.rgb = RGBColor(112, 128, 144) # Slate Gray
    
    # Insert Banner Image on Cover Page
    if os.path.exists("assets/travelmate_banner.png"):
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_before = Pt(20)
        p_banner.paragraph_format.space_after = Pt(20)
        run_banner = p_banner.add_run()
        run_banner.add_picture("assets/travelmate_banner.png", width=Inches(4.5))
        
    p_submitted_by = doc.add_paragraph()
    p_submitted_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_submitted_by.paragraph_format.space_before = Pt(24)
    r_sub_label = p_submitted_by.add_run('Submitted by\n')
    r_sub_label.font.size = Pt(11)
    r_sub_label.italic = True
    
    r_student_name = p_submitted_by.add_run('Gangavaram Joshika\n')
    r_student_name.font.size = Pt(13)
    r_student_name.bold = True
    r_student_name.font.color.rgb = RGBColor(27, 54, 93)
    
    r_student_roll = p_submitted_by.add_run('24STUCHH010445\n\n')
    r_student_roll.font.size = Pt(11)
    r_student_roll.bold = True
    
    r_super_label = p_submitted_by.add_run('Under the Supervision of\n')
    r_super_label.font.size = Pt(11)
    r_super_label.italic = True
    
    r_super_name = p_submitted_by.add_run('Dr V V SATYANARAYANA MURTHY B\n')
    r_super_name.font.size = Pt(12)
    r_super_name.bold = True
    r_super_name.font.color.rgb = RGBColor(27, 54, 93)
    
    r_super_desig = p_submitted_by.add_run('As part of the Internship Program - I\n')
    r_super_desig.font.size = Pt(11)
    
    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.space_before = Pt(36)
    
    r_dept = p_dept.add_run(
        'Department of Artificial Intelligence And Machine Learning\n'
        'ICFAI Foundation for Higher Education, Hyderabad, India\n'
    )
    r_dept.font.size = Pt(11)
    r_dept.bold = True
    
    r_date = p_dept.add_run('June 2026')
    r_date.font.size = Pt(11)
    r_date.bold = True
    r_date.font.color.rgb = RGBColor(112, 128, 144)
    
    doc.add_page_break()
    
    # Configure Headers & Footers for the rest of the document
    # Note: Word allows different first page, so the cover page won't have headers/footers
    body_section = doc.sections[-1]
    body_section.different_first_page_header_footer = True
    
    # Footer setup
    footer = body_section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run_text = f_p.add_run("TravelMate AI Project Report | Page ")
    f_run_text.font.size = Pt(9)
    f_run_text.font.color.rgb = RGBColor(128, 128, 128)
    f_run_num = f_p.add_run()
    f_run_num.font.size = Pt(9)
    f_run_num.font.color.rgb = RGBColor(128, 128, 128)
    f_run_num.bold = True
    add_page_number(f_run_num)
    
    # Header setup
    header = body_section.header
    h_p = header.paragraphs[0]
    h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h_run = h_p.add_run("Department of AI & ML, IcfaiTech")
    h_run.font.size = Pt(9)
    h_run.font.color.rgb = RGBColor(128, 128, 128)
    h_run.italic = True
    
    # ==========================================
    # 2. CERTIFICATE
    # ==========================================
    add_custom_heading(doc, "Certificate of Approval", level=1, space_before=24, space_after=18)
    
    p_cert_body = (
        "This is to certify that the project report titled \"TRAVELMATE AI: SMART TRAVEL COMPANION AND "
        "GUIDANCE PLATFORM\" is a bonafide record of work carried out by Gangavaram Joshika (Roll No: "
        "24STUCHH010445) under my supervision and guidance. The work described in this report is "
        "submitted in partial fulfillment of the requirements for the award of the Internship-1 for the "
        "degree of Bachelor of Technology in Artificial Intelligence and Machine Learning from ICFAI Foundation "
        "for Higher Education, Hyderabad, and has not been submitted elsewhere for any other degree or diploma."
    )
    add_custom_paragraph(doc, p_cert_body, space_after=48)
    
    # Signature Lines
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Clear borders
    tblPr = table_sig._tbl.tblPr
    borders = parse_xml(r'<w:tblBorders {} ><w:top w:val="none"/><w:bottom w:val="none"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>'.format(nsdecls('w')))
    tblPr.append(borders)
    
    # Add spacing and text
    row_sig1 = table_sig.rows[0]
    row_sig2 = table_sig.rows[1]
    
    row_sig1.cells[0].paragraphs[0].text = "______________________________\nInternal Supervisor"
    row_sig1.cells[1].paragraphs[0].text = "______________________________\nHead of Department"
    row_sig2.cells[0].paragraphs[0].text = "Dr V V SATYANARAYANA MURTHY B\nDepartment of AI & ML"
    row_sig2.cells[1].paragraphs[0].text = "Faculty of Science & Technology\nICFAI Foundation for Higher Education"
    
    for row in table_sig.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    
    doc.add_page_break()
    
    # ==========================================
    # 3. ACKNOWLEDGEMENT
    # ==========================================
    add_custom_heading(doc, "Acknowledgement", level=1, space_before=24, space_after=18)
    
    p_ack1 = (
        "I express my deepest gratitude to my esteemed project supervisor, Dr V V SATYANARAYANA MURTHY B, "
        "Department of Artificial Intelligence and Machine Learning, for his invaluable guidance, continuous "
        "encouragement, and constructive feedback throughout the course of this project. His expertise and "
        "insightful suggestions have played a critical role in shaping the architecture and implementation "
        "of this platform."
    )
    add_custom_paragraph(doc, p_ack1, space_after=12)
    
    p_ack2 = (
        "I would also like to thank the Head of the Department and the administration of IcfaiTech, Faculty of "
        "Science & Technology, ICFAI Foundation for Higher Education, Hyderabad, for providing the necessary "
        "academic resources, infrastructure, and an encouraging environment that enabled the successful execution "
        "of this internship project."
    )
    add_custom_paragraph(doc, p_ack2, space_after=12)
    
    p_ack3 = (
        "Finally, I extend my heartfelt appreciation to my family and peers for their constant support, patience, "
        "and motivation, which helped me stay focused and overcome technical challenges during the development "
        "of TravelMate AI."
    )
    add_custom_paragraph(doc, p_ack3, space_after=24)
    
    p_stud_sig = doc.add_paragraph()
    p_stud_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_stud_sig = p_stud_sig.add_run("Gangavaram Joshika\nRoll No: 24STUCHH010445")
    r_stud_sig.font.name = 'Calibri'
    r_stud_sig.font.size = Pt(11)
    r_stud_sig.bold = True
    
    doc.add_page_break()
    
    # ==========================================
    # 4. ABSTRACT
    # ==========================================
    add_custom_heading(doc, "Abstract", level=1, space_before=24, space_after=18)
    
    p_abs1 = (
        "Modern tourism relies heavily on digital information, yet travelers frequently face the challenge of "
        "fragmented data across multiple websites when planning a trip. Information regarding regulatory checklists, "
        "visa requirements, local traffic and criminal laws, cultural etiquette, local cuisines, and accommodation "
        "options must be compiled manually. Furthermore, budget-conscious travelers require customized, structured "
        "itineraries, while international tourists benefit from context-aware, real-time, multilingual guidance."
    )
    add_custom_paragraph(doc, p_abs1, space_after=12)
    
    p_abs2 = (
        "To address these challenges, this project presents TravelMate AI, a comprehensive, Streamlit-based web "
        "application designed to serve as an intelligent, all-in-one travel companion. The platform integrates a "
        "local SQLite database pre-seeded with regulatory, geographical, and cultural profiles for multiple "
        "destinations. The frontend is built using Streamlit, styled with custom CSS injections and Outfit typography. "
        "Key modules include a Home Dashboard with global search, a Country Guide displaying pre-departure checklists "
        "and cultural dos/don'ts, a City Explorer showcasing rated attractions, cuisines, and hotel tiers, and a "
        "Smart Travel Planner that dynamically generates day-by-day itineraries (1-7 days) paired with interactive "
        "Plotly budget charts."
    )
    add_custom_paragraph(doc, p_abs2, space_after=12)
    
    p_abs3 = (
        "A core innovation of the platform is the context-aware AI Travel Assistant, a conversational chatbot. "
        "Using the Google Agent Development Kit (google-adk) and the gemini-2.5-flash model, the chatbot autonomously "
        "executes SQL database tool queries. When API access is unavailable, the system gracefully degrades to an "
        "offline, rule-based keyword matching parser that utilizes active UI session contexts (selected country/city) "
        "to resolve ambiguous queries. Security is maintained through JWT-based session tokens and SHA-256 password "
        "hashing. The codebase enforces strict quality standards via automated Pytest suites and pre-commit hooks "
        "covering formatting, security (Bandit), and static type checks (Mypy). TravelMate AI represents a scalable, "
        "local-first solution that enhances traveler safety, cultural compliance, and planning efficiency."
    )
    add_custom_paragraph(doc, p_abs3, space_after=24)
    
    doc.add_page_break()
    
    # ==========================================
    # 5. TABLE OF CONTENTS
    # ==========================================
    add_custom_heading(doc, "Table of Contents", level=1, space_before=24, space_after=18)
    
    toc_data = [
        ("Abstract", "iii"),
        ("1. Introduction", "1"),
        ("   1.1 Project Overview", "1"),
        ("   1.2 Problem Statement", "1"),
        ("   1.3 Objectives", "2"),
        ("2. Literature Survey & Gap Analysis", "3"),
        ("   2.1 Existing Travel Guidance Systems", "3"),
        ("   2.2 Comparison of Travel Platforms", "3"),
        ("   2.3 Research Gap", "4"),
        ("3. System Design and Architecture", "5"),
        ("   3.1 Architectural Overview", "5"),
        ("   3.2 System Workflow Flowchart", "6"),
        ("   3.3 Advantages of the Proposed System", "6"),
        ("4. Database Design", "7"),
        ("   4.1 Relational Database Schema", "7"),
        ("   4.2 Detailed Table Structures", "7"),
        ("5. Modules Description", "10"),
        ("   5.1 Home Dashboard & Search", "10"),
        ("   5.2 Country Information Guide", "10"),
        ("   5.3 City Explorer", "10"),
        ("   5.4 Smart Travel Itinerary Planner", "11"),
        ("   5.5 AI Travel Assistant Chatbot", "11"),
        ("   5.6 User Authentication & Profile Persistence", "11"),
        ("6. Implementation Details", "12"),
        ("   6.1 Key Code Structure", "12"),
        ("   6.2 Main Entry and Navigation", "12"),
        ("   6.3 JWT Authentication and Session Management", "13"),
        ("7. Algorithms & AI Integration", "14"),
        ("   7.1 Autonomous AI Agent (Google ADK)", "14"),
        ("   7.2 Rule-Based Context-Aware Fallback", "15"),
        ("8. Application Interface Screenshots", "16"),
        ("9. Testing & Quality Assurance", "19"),
        ("   9.1 BDD-style Pytest Suite", "19"),
        ("   9.2 Quality Assurance and Pre-commit Hooks", "20"),
        ("10. Advantages, Limitations, & Future Scope", "21"),
        ("   10.1 Advantages", "21"),
        ("   10.2 Limitations", "21"),
        ("   10.3 Future Scope", "21"),
        ("11. Conclusion", "23"),
        ("References", "24"),
    ]
    
    table_toc = create_styled_table(doc, len(toc_data) + 1, 2)
    style_table_header(table_toc.rows[0], ["Section Title", "Page Number"])
    
    for i, (title, page) in enumerate(toc_data):
        style_table_row(table_toc.rows[i + 1], [title, page], is_even=(i % 2 == 1))
        
    doc.add_page_break()
    
    # ==========================================
    # 6. INTRODUCTION & OBJECTIVES
    # ==========================================
    add_custom_heading(doc, "1. Introduction", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "1.1 Project Overview", level=2, space_before=12, space_after=6)
    p_intro1 = (
        "In an increasingly globalized world, tourism has evolved from a luxury into a highly accessible "
        "activity. However, planning international or domestic travel remains a complex and fragmented "
        "process. Travelers must research visa requirements, local emergency services, transit systems, "
        "accommodation options, and dining details. Beyond logistics, a significant challenge is understanding "
        "local laws and cultural etiquette, which vary drastically across countries. For instance, regulations "
        "regarding jaywalking, littering, or public consumption of certain items carry heavy fines in places "
        "like Singapore, while cultural norms regarding tipping or dress codes dictate social interactions in Japan."
    )
    add_custom_paragraph(doc, p_intro1)
    
    p_intro2 = (
        "TravelMate AI is a comprehensive, local-first web application built on the Streamlit framework, designed "
        "to serve as an intelligent, all-in-one travel companion. By combining structured relational data with "
        "context-aware artificial intelligence, TravelMate AI consolidates visa guidelines, regulatory warnings, "
        "cultural etiquette, local transit guides, and hotel listings into a single responsive interface. The platform "
        "also features a dynamic, budget-oriented travel planner and a multi-lingual AI chatbot that acts as an "
        "autonomous database agent."
    )
    add_custom_paragraph(doc, p_intro2)
    
    add_custom_heading(doc, "1.2 Problem Statement", level=2, space_before=12, space_after=6)
    p_prob = (
        "Traditional travel planning is highly fragmented, requiring tourists to visit multiple independent web portals "
        "to compile an itinerary, verify visa rules, list local laws, and locate emergency numbers. This manual compilation "
        "is time-consuming and prone to errors. Existing travel platforms (such as TripAdvisor and Google Travel) excel "
        "at business recommendations but fail to provide centralized regulatory, legal, and safety checklists. "
        "Additionally, static itinerary planning does not adapt dynamically to custom durations or budget tiers, "
        "and most platforms lack context-aware, real-time conversational assistants that operate locally. There is a clear "
        "need for a unified, secure, and multi-lingual system that provides structured destination profiles alongside "
        "intelligent, context-sensitive chat assistance."
    )
    add_custom_paragraph(doc, p_prob)
    
    add_custom_heading(doc, "1.3 Objectives", level=2, space_before=12, space_after=6)
    p_obj_intro = "The primary objectives of the TravelMate AI project are:"
    add_custom_paragraph(doc, p_obj_intro)
    
    add_bullet_point(doc, "1. Centralize Travel Information: ", "To design and implement a relational SQLite database storing regulatory, visa, safety, and cultural guidelines alongside city-specific attractions, hotels, dining, and transit options.")
    add_bullet_point(doc, "2. Provide Dynamic Itinerary Planning: ", "To develop a customizable planning module that generates day-by-day timelines for 1-7 days, paired with interactive Plotly budget visualizations.")
    add_bullet_point(doc, "3. Implement Context-Aware Conversational AI: ", "To integrate an autonomous database agent using the Google Agent Development Kit (google-adk) and Gemini models, with a robust rule-based keyword fallback that utilizes UI session states.")
    add_bullet_point(doc, "4. Support Multi-lingual Accessibility: ", "To build internationalization (i18n) support, offering complete UI and database translation in English, Hindi, and Telugu.")
    add_bullet_point(doc, "5. Secure User Data: ", "To establish JWT-based session management and SHA-256 password hashing for user registration, profile editing, and travel history persistence.")
    add_bullet_point(doc, "6. Enforce Software Quality: ", "To implement a comprehensive BDD-style test suite using Pytest and integrate pre-commit hooks for code style, static analysis, and security scanning.")
    
    doc.add_page_break()
    
    # ==========================================
    # 7. LITERATURE SURVEY & GAP ANALYSIS
    # ==========================================
    add_custom_heading(doc, "2. Literature Survey & Gap Analysis", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "2.1 Existing Travel Guidance Systems", level=2, space_before=12, space_after=6)
    p_lit1 = (
        "Digital travel assistance has seen significant growth, with platforms categorized into commercial review portals, "
        "governmental advisory sites, and automated itinerary generators. Commercial platforms like TripAdvisor and Yelp "
        "rely on user-generated content, offering excellent restaurant and hotel reviews. However, they lack structured "
        "legal and cultural guidance. Government portals (such as the US Department of State Travel Advisories) provide "
        "excellent safety and visa updates but are text-heavy, lack local city guides, and offer no planning utilities. "
        "Automated planning apps (like Sygic Travel) offer day-by-day mapping but require paid subscriptions, lack local "
        "regulatory context, and do not integrate conversational AI agents."
    )
    add_custom_paragraph(doc, p_lit1)
    
    add_custom_heading(doc, "2.2 Comparison of Travel Platforms", level=2, space_before=12, space_after=6)
    p_table_desc = "The following table compares the features of TravelMate AI with major existing digital travel platforms:"
    add_custom_paragraph(doc, p_table_desc)
    
    table_comp = create_styled_table(doc, 5, 5)
    style_table_header(table_comp.rows[0], ["Feature / Metric", "TripAdvisor", "Google Travel", "Govt. Advisories", "TravelMate AI (Proposed)"])
    
    comp_data = [
        ("Centralized Visa & Regulatory Info", "No (Fragmented)", "No", "Yes (Text-only)", "Yes (Structured Table)"),
        ("Local Law & Etiquette Alerts", "No", "No", "Yes (General)", "Yes (Specific / Detailed)"),
        ("Dynamic Budget Visualizer", "No", "No (Flights/Hotels only)", "No", "Yes (Interactive Plotly Chart)"),
        ("Context-Aware AI Chatbot", "No", "Yes (Gemini - General)", "No", "Yes (Autonomous DB Agent & Fallback)"),
    ]
    
    for idx, row_vals in enumerate(comp_data):
        style_table_row(table_comp.rows[idx + 1], row_vals, is_even=(idx % 2 == 1))
        
    add_custom_heading(doc, "2.3 Research Gap", level=2, space_before=12, space_after=6)
    p_gap = (
        "A critical analysis of existing systems reveals a major research gap: the lack of integration between "
        "commercial tourism listings and regulatory/cultural safety guidelines. Tourists often travel without knowing "
        "local restrictions or cultural norms, leading to legal penalties or social discomfort. Furthermore, standard "
        "itinerary planners do not provide instant, context-sensitive chat support linked to the user's active search "
        "context. TravelMate AI bridges this gap by combining regulatory country profiles, localized city explorer tables, "
        "dynamic budgeting, and an autonomous AI chatbot that inherits the active UI selection as its conversational context."
    )
    add_custom_paragraph(doc, p_gap)
    
    doc.add_page_break()
    
    # ==========================================
    # 8. SYSTEM DESIGN AND ARCHITECTURE
    # ==========================================
    add_custom_heading(doc, "3. System Design and Architecture", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "3.1 Architectural Overview", level=2, space_before=12, space_after=6)
    p_arch = (
        "TravelMate AI implements a decoupled three-layer architecture to ensure modularity, scalability, and ease of maintenance. "
        "The system separates user interaction, business/agent logic, and data storage into distinct layers:"
    )
    add_custom_paragraph(doc, p_arch)
    
    add_bullet_point(doc, "1. Presentation Layer (Frontend): ", "Built using Streamlit. It consists of multiple view files (home, country_info, city_info, planner, chatbot, profile, history, saved_trips, auth) that render dynamic UI components, charts, and input controls. Custom CSS is injected to enforce the 'Outfit' typography and a modern dark/light card-based aesthetic.")
    add_bullet_point(doc, "2. Application / Logic Layer (Backend): ", "Handles the core processing, including JWT-based session management, password cryptography (SHA-256), multi-lingual translation dictionary lookups, and the AI chatbot orchestration. The chatbot module integrates the Google Agent Development Kit (google-adk) to coordinate autonomous tool calls or execute the rule-based keyword matching algorithm.")
    add_bullet_point(doc, "3. Data Access Layer (Database): ", "Consists of an embedded SQLite database (travel.db) managed via Python's sqlite3 library. This layer executes connection pooling, parameterized SQL queries, and transactional updates for user profiles, search histories, and saved itineraries.")
    
    add_custom_heading(doc, "3.2 System Workflow Flowchart", level=2, space_before=12, space_after=6)
    p_flow_desc = "The flowchart below describes the step-by-step data and logic flow when a user interacts with the platform:"
    add_custom_paragraph(doc, p_flow_desc)
    
    # Text-based flowchart
    flowchart_text = (
        "  [User Input / Query] \n"
        "           │\n"
        "           ▼\n"
        "  [Streamlit Frontend (app.py / pages)] ──(Renders UI & Tracks Session State)\n"
        "           │\n"
        "           ├──────────────────────────────┐\n"
        "           ▼                              ▼\n"
        "  [Authentication (JWT)]        [AI Chatbot Agent (chatbot.py)]\n"
        "           │                              │\n"
        "           │                              ├───────────────┐ (If API Key Present)\n"
        "           │                              ▼               ▼\n"
        "           │                    [Google ADK Engine]   [Rule-based Fallback]\n"
        "           │                    (Autonomous Tools)    (Keyword Matching)\n"
        "           │                              │               │\n"
        "           └──────────────┬───────────────┴───────────────┘\n"
        "                          │\n"
        "                          ▼\n"
        "             [Database Utility (database.py)]\n"
        "                          │\n"
        "                          ▼\n"
        "             [SQLite Database (travel.db)]\n"
    )
    p_flow = doc.add_paragraph()
    p_flow.paragraph_format.space_before = Pt(6)
    p_flow.paragraph_format.space_after = Pt(12)
    p_flow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_flow = p_flow.add_run(flowchart_text)
    run_flow.font.name = 'Courier New'
    run_flow.font.size = Pt(9.5)
    run_flow.font.color.rgb = RGBColor(27, 54, 93)
    
    add_custom_heading(doc, "3.3 Advantages of the Proposed System", level=2, space_before=12, space_after=6)
    add_bullet_point(doc, "- Consolidated Information: ", "Combines logistical, regulatory, and cultural data in a single dashboard.")
    add_bullet_point(doc, "- Dynamic Planning & Visualization: ", "Generates detailed, day-by-day schedules with interactive budget breakdowns.")
    add_bullet_point(doc, "- Context-Aware Conversational Interface: ", "The AI assistant automatically inherits the active UI search context, minimizing user input.")
    add_bullet_point(doc, "- Multi-lingual Localization: ", "Complete localization in English, Hindi, and Telugu expands accessibility.")
    add_bullet_point(doc, "- Local-First & Highly Secure: ", "An embedded SQLite database requires zero configuration, while JWTs and password hashing ensure user data privacy.")
    
    doc.add_page_break()
    
    # ==========================================
    # 9. DATABASE DESIGN
    # ==========================================
    add_custom_heading(doc, "4. Database Design", level=1, space_before=24, space_after=12)
    
    p_db_intro = (
        "The relational database schema is implemented in SQLite (travel.db). It is designed to maintain "
        "strict data integrity and support fast query execution. The database consists of six tables: "
        "countries, cities, users, travel_history, saved_trips, and weather_history. The schema details "
        "for each table are described below:"
    )
    add_custom_paragraph(doc, p_db_intro)
    
    # 1. countries Table
    add_custom_heading(doc, "4.1 Table: countries", level=2, space_before=12, space_after=6)
    table_countries = create_styled_table(doc, 12, 3)
    style_table_header(table_countries.rows[0], ["Column Name", "Data Type", "Description"])
    countries_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique country identifier, auto-incremented"),
        ("country_name", "TEXT UNIQUE", "Name of the country (e.g., India, Japan, Singapore)"),
        ("capital", "TEXT", "Capital city of the country"),
        ("currency", "TEXT", "Local currency symbol and code (e.g., INR, JPY, SGD)"),
        ("language", "TEXT", "Primary official language(s) spoken"),
        ("timezone", "TEXT", "Time zone offset (e.g., GMT+9, GMT+8)"),
        ("emergency_number", "TEXT", "Emergency contact numbers (police, medical, fire)"),
        ("visa_info", "TEXT", "Visa requirements and pre-departure checklists"),
        ("rules", "TEXT", "Crucial local laws and regulatory warnings"),
        ("etiquette", "TEXT", "Cultural dos and don'ts for travelers"),
        ("safety_tips", "TEXT", "General safety recommendations and crime warnings"),
    ]
    for idx, row in enumerate(countries_schema):
        style_table_row(table_countries.rows[idx + 1], row, is_even=(idx % 2 == 1))
        
    # 2. cities Table
    add_custom_heading(doc, "4.2 Table: cities", level=2, space_before=12, space_after=6)
    table_cities = create_styled_table(doc, 12, 3)
    style_table_header(table_cities.rows[0], ["Column Name", "Data Type", "Description"])
    cities_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique city identifier, auto-incremented"),
        ("country_id", "INTEGER", "Foreign Key referencing countries(id) with Cascade Delete"),
        ("city_name", "TEXT", "Name of the city (e.g., Tokyo, Osaka, Hyderabad)"),
        ("description", "TEXT", "Brief geographical and historical overview"),
        ("transport_info", "TEXT", "Local transit details (buses, trains, subways)"),
        ("food_info", "TEXT (JSON)", "JSON array of local delicacies and descriptions"),
        ("tourist_places", "TEXT (JSON)", "JSON array of attractions, ratings, and best visit times"),
        ("hotel_info", "TEXT (JSON)", "JSON object of hotels categorized by price tier"),
        ("shopping_areas", "TEXT", "Popular local markets, shopping streets, and malls"),
        ("airport_details", "TEXT", "Local airport names and transfer guidelines"),
        ("safety_recommendations", "TEXT", "Neighborhood-specific safety advisories"),
    ]
    for idx, row in enumerate(cities_schema):
        style_table_row(table_cities.rows[idx + 1], row, is_even=(idx % 2 == 1))
        
    # 3. users Table
    add_custom_heading(doc, "4.3 Table: users", level=2, space_before=12, space_after=6)
    table_users = create_styled_table(doc, 10, 3)
    style_table_header(table_users.rows[0], ["Column Name", "Data Type", "Description"])
    users_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique user identifier, auto-incremented"),
        ("full_name", "TEXT", "Full name of the registered user"),
        ("email", "TEXT UNIQUE", "Unique email address used for login"),
        ("phone", "TEXT", "Contact phone number"),
        ("country", "TEXT", "User's country of residence"),
        ("city", "TEXT", "User's city of residence"),
        ("profile_pic", "TEXT", "Base64 encoded string of profile picture or image URL"),
        ("password_hash", "TEXT", "SHA-256 hashed password string"),
        ("preferences", "TEXT (JSON)", "JSON array of user travel preferences (e.g., vegetarian, luxury)"),
    ]
    for idx, row in enumerate(users_schema):
        style_table_row(table_users.rows[idx + 1], row, is_even=(idx % 2 == 1))

    # 4. travel_history Table
    add_custom_heading(doc, "4.4 Table: travel_history", level=2, space_before=12, space_after=6)
    table_history = create_styled_table(doc, 7, 3)
    style_table_header(table_history.rows[0], ["Column Name", "Data Type", "Description"])
    history_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique history record identifier"),
        ("user_id", "INTEGER", "Foreign Key referencing users(id)"),
        ("activity_type", "TEXT", "Type of activity ('search', 'itinerary', 'chat')"),
        ("query", "TEXT", "User input query or destination searched"),
        ("details", "TEXT (JSON)", "JSON string containing metadata or generated itinerary details"),
        ("is_favorite", "INTEGER", "Binary flag (0 or 1) indicating if the item is bookmarked"),
    ]
    for idx, row in enumerate(history_schema):
        style_table_row(table_history.rows[idx + 1], row, is_even=(idx % 2 == 1))

    # 5. saved_trips Table
    add_custom_heading(doc, "4.5 Table: saved_trips", level=2, space_before=12, space_after=6)
    table_saved = create_styled_table(doc, 8, 3)
    style_table_header(table_saved.rows[0], ["Column Name", "Data Type", "Description"])
    saved_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique saved trip identifier"),
        ("user_id", "INTEGER", "Foreign Key referencing users(id)"),
        ("trip_type", "TEXT", "Type of saved item ('itinerary', 'destination', 'hotel')"),
        ("name", "TEXT", "Name of the saved trip or destination"),
        ("collection_name", "TEXT", "Name of the folder/collection (defaults to 'My Saved Trips')"),
        ("details", "TEXT", "JSON string containing complete details of the saved item"),
        ("travel_date", "TEXT", "Planned travel date in YYYY-MM-DD format"),
    ]
    for idx, row in enumerate(saved_schema):
        style_table_row(table_saved.rows[idx + 1], row, is_even=(idx % 2 == 1))

    # 6. weather_history Table
    add_custom_heading(doc, "4.6 Table: weather_history", level=2, space_before=12, space_after=6)
    table_weather = create_styled_table(doc, 8, 3)
    style_table_header(table_weather.rows[0], ["Column Name", "Data Type", "Description"])
    weather_schema = [
        ("id", "INTEGER PRIMARY KEY", "Unique weather record identifier"),
        ("city_id", "INTEGER", "Foreign Key referencing cities(id)"),
        ("month_num", "INTEGER", "Month number (1 = January, 12 = December)"),
        ("month_name", "TEXT", "Full name of the month (e.g., January)"),
        ("avg_temp", "REAL", "Average monthly temperature in degrees Celsius"),
        ("rainfall", "REAL", "Average monthly rainfall in millimeters"),
        ("description", "TEXT", "General weather description (e.g., Hot & Humid)"),
    ]
    for idx, row in enumerate(weather_schema):
        style_table_row(table_weather.rows[idx + 1], row, is_even=(idx % 2 == 1))

    doc.add_page_break()
    
    # ==========================================
    # 10. MODULES DESCRIPTION
    # ==========================================
    add_custom_heading(doc, "5. Modules Description", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "5.1 Home Dashboard & Search", level=2, space_before=12, space_after=6)
    p_mod_home = (
        "The Home Dashboard serves as the primary landing page. It features a modern hero header and "
        "showcases featured destination cards (such as Hyderabad, Singapore, and Tokyo). A global search "
        "bar allows users to type queries like 'shrine', 'biryani', or country/city names. The search engine "
        "queries multiple columns across both countries and cities tables, returning matching destinations with "
        "direct links to their profiles. A sidebar dropdown allows users to select an active country and city, "
        "which sets the global context across the application."
    )
    add_custom_paragraph(doc, p_mod_home)
    
    add_custom_heading(doc, "5.2 Country Information Guide", level=2, space_before=12, space_after=6)
    p_mod_country = (
        "The Country Guide provides regulatory, legal, and cultural information. It displays a profile summary "
        "including the capital, currency, languages, timezone, and local emergency contact numbers. A key feature "
        "is the Pre-departure Checklist, which lists visa requirements and essential travel items. The module "
        "also contains a 'Rules & Etiquette' section detailing critical local laws (e.g., bans on chewing gum in "
        "Singapore or traffic regulations in Japan) and cultural dos and don'ts to ensure respectful travel."
    )
    add_custom_paragraph(doc, p_mod_country)
    
    add_custom_heading(doc, "5.3 City Explorer", level=2, space_before=12, space_after=6)
    p_mod_city = (
        "The City Explorer offers localized tourist recommendations. It is divided into four sections: "
        "(1) Sightseeing & Attractions: displays popular spots, their star ratings, and the best time of day to "
        "visit; (2) Stays & Accommodations: lists hotels categorized into Budget, Mid-range, and Luxury tiers with "
        "estimated prices; (3) Cuisines & Dining: showcases traditional vegetarian and non-vegetarian dishes; "
        "(4) Transit Guidelines: provides local transportation details and airport transfer options."
    )
    add_custom_paragraph(doc, p_mod_city)
    
    add_custom_heading(doc, "5.4 Smart Travel Itinerary Planner", level=2, space_before=12, space_after=6)
    p_mod_planner = (
        "The Itinerary Planner allows users to customize their travel schedules. Users select a trip duration "
        "(1-7 days) and a budget level (Economy, Mid-range, Luxury). The system dynamically generates a day-by-day "
        "timeline featuring morning, lunch, afternoon, and evening activities. Simultaneously, it renders an "
        "interactive Plotly Pie Chart representing the cost allocation across accommodation, dining, transit, and "
        "shopping, accompanied by an itemized budget table. Registered users can save these itineraries to their profile."
    )
    add_custom_paragraph(doc, p_mod_planner)
    
    add_custom_heading(doc, "5.5 AI Travel Assistant Chatbot", level=2, space_before=12, space_after=6)
    p_mod_chat = (
        "The AI Chatbot provides real-time conversational assistance. It operates in two modes: "
        "(1) Agent Kit Mode: if a Gemini API key is configured, it utilizes the Google Agent Development Kit "
        "to run an autonomous agent that calls database query tools; (2) Offline Mode: if the API key is absent, "
        "it falls back to a local rule-based keyword matching parser. The chatbot is context-aware: if a user "
        "asks a general question like 'where should I stay?', the bot automatically retrieves the active country "
        "or city selected in the sidebar session state."
    )
    add_custom_paragraph(doc, p_mod_chat)
    
    add_custom_heading(doc, "5.6 User Authentication & Profile", level=2, space_before=12, space_after=6)
    p_mod_auth = (
        "The Authentication module manages user accounts. Users can register, log in, and secure their sessions "
        "using a 'Remember Me' token. The system generates a JWT-like token using HMAC-SHA256 and hashes passwords "
        "locally using SHA-256 with a static salt. Once logged in, users can edit their profile details, view "
        "dashboard analytics (total trips planned, countries explored, and AI usage stats), view their search history, "
        "and access saved itineraries."
    )
    add_custom_paragraph(doc, p_mod_auth)
    
    doc.add_page_break()
    
    # ==========================================
    # 11. IMPLEMENTATION DETAILS
    # ==========================================
    add_custom_heading(doc, "6. Implementation Details", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "6.1 Key Code Structure", level=2, space_before=12, space_after=6)
    p_imp_struct = (
        "The codebase is structured to enforce separation of concerns. The entry point is app.py, which "
        "configures the multi-page router. View components are stored in the pages/ directory, database "
        "operations are defined in utils/database.py, cryptographic functions are in utils/auth_utils.py, "
        "and global CSS injections are managed by utils/styles.py. Internationalization (i18n) is handled "
        "by utils/i18n.py, which stores translation mappings for English, Hindi, and Telugu UI elements."
    )
    add_custom_paragraph(doc, p_imp_struct)
    
    add_custom_heading(doc, "6.2 Main Entry and Navigation", level=2, space_before=12, space_after=6)
    p_imp_nav = (
        "The navigation and routing are implemented using Streamlit's st.Page and st.navigation APIs. "
        "The sidebar dynamically updates the available pages based on the user's login status. "
        "If a user is logged in, their profile, travel history, and saved trips are added under an 'Account' "
        "group, and their name, email, and profile picture are displayed in a styled header card. If the user "
        "is a guest, an 'Auth' page is displayed instead."
    )
    add_custom_paragraph(doc, p_imp_nav)
    
    add_custom_heading(doc, "6.3 JWT Authentication and Session Management", level=2, space_before=12, space_after=6)
    p_imp_jwt = (
        "Password security is implemented using SHA-256 hashing with a static salt to prevent rainbow table attacks. "
        "Session persistence is achieved by generating a signed JWT-like token containing the user's ID, name, email, "
        "and an expiration timestamp. The token is signed using HMAC-SHA256 with a secret key. When the application "
        "starts, it checks the session state for a 'remember_me_token'. If present and valid, it decodes the token "
        "and automatically logs the user in, restoring their session without requiring re-authentication."
    )
    add_custom_paragraph(doc, p_imp_jwt)
    
    doc.add_page_break()
    
    # ==========================================
    # 12. ALGORITHMS & AI INTEGRATION
    # ==========================================
    add_custom_heading(doc, "7. Algorithms & AI Integration", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "7.1 Autonomous AI Agent (Google ADK)", level=2, space_before=12, space_after=6)
    p_alg_adk = (
        "The AI Travel Assistant integrates the Google Agent Development Kit (google-adk) and the "
        "gemini-2.5-flash model. The agent is defined with a natural language instruction set and is equipped "
        "with three python functions registered as tools: get_country_info, get_city_info, and search_destinations. "
        "When the user submits a chat query, the Agent Runner analyzes the request and autonomously decides "
        "which tool to execute, parses the tool's output, and formats a coherent markdown response."
    )
    add_custom_paragraph(doc, p_alg_adk)
    
    # Show code block
    p_code_adk = doc.add_paragraph()
    p_code_adk.paragraph_format.space_before = Pt(6)
    p_code_adk.paragraph_format.space_after = Pt(12)
    run_code_adk = p_code_adk.add_run(
        "travel_agent = Agent(\n"
        "    name=\"travel_mate_agent\",\n"
        "    model=\"gemini-2.5-flash\",\n"
        "    instruction=\"You are TravelMate AI. Use the database tools "
        "provided (get_country_info, get_city_info, search_destinations) "
        "to answer user queries. If the query does not mention a location, "
        "refer to the Active Context in the system message.\",\n"
        "    tools=[get_country_info, get_city_info, search_destinations]\n"
        ")"
    )
    run_code_adk.font.name = 'Courier New'
    run_code_adk.font.size = Pt(9)
    run_code_adk.font.color.rgb = RGBColor(27, 54, 93)
    
    add_custom_heading(doc, "7.2 Rule-Based Context-Aware Fallback", level=2, space_before=12, space_after=6)
    p_alg_fall = (
        "To ensure the application remains functional offline or when an API key is not configured, a "
        "local rule-based fallback algorithm is implemented in pages/chatbot.py. The algorithm operates as follows:"
    )
    add_custom_paragraph(doc, p_alg_fall)
    
    add_bullet_point(doc, "1. Tokenization & Normalization: ", "The user's query is normalized (converted to English using i18n helpers) and tokenized into lowercase words.")
    add_bullet_point(doc, "2. Entity Recognition: ", "The tokens are matched against country and city names in the database. If a match is found, that entity is set as the target. If no match is found, the system retrieves the active country/city from the UI session state (context fallback).")
    add_bullet_point(doc, "3. Topic Detection: ", "The query is checked for specific keyword sets. For example, keywords like 'food', 'cuisine', or 'biryani' trigger the 'Food' topic; 'rules', 'law', or 'chewing gum' trigger the 'Rules' topic; 'hotel' or 'stay' trigger the 'Accommodations' topic.")
    add_bullet_point(doc, "4. Database Querying & Formatting: ", "Based on the detected topic and target entity, the system executes the corresponding SQL query, parses the JSON payload (if applicable, such as food_info or tourist_places), and formats a structured response in the user's selected language.")
    
    doc.add_page_break()
    
    # ==========================================
    # 13. APPLICATION INTERFACE SCREENSHOTS
    # ==========================================
    add_custom_heading(doc, "8. Application Interface Screenshots", level=1, space_before=24, space_after=12)
    
    p_scr_intro = (
        "This section contains screenshots of the working TravelMate AI application, demonstrating "
        "its user interface, multi-page layout, and features."
    )
    add_custom_paragraph(doc, p_scr_intro)
    
    # Figure 1: Banner
    add_figure(doc, "assets/travelmate_banner.png", "Figure 1: TravelMate AI Application Banner and Branding", width_in_inches=5.0)
    p_desc_fig1 = (
        "Figure 1 shows the branding banner for TravelMate AI. The design features a modern, clean "
        "aesthetic with the slogan 'Smart Travel Companion'. It is used at the top of the Home page and "
        "on the cover page to establish a premium visual identity."
    )
    add_custom_paragraph(doc, p_desc_fig1)
    
    # Figure 2: Hyderabad City
    add_figure(doc, "assets/hyderabad_city.png", "Figure 2: City Explorer - Hyderabad Page", width_in_inches=5.0)
    p_desc_fig2 = (
        "Figure 2 illustrates the City Explorer page when 'Hyderabad' is selected. The interface displays "
        "detailed local profiles including sightseeing spots (such as Charminar and Golconda Fort) with star "
        "ratings and recommended times to visit. It also lists local cuisines (e.g., Hyderabadi Biryani, Double ka Meetha) "
        "and categorizes accommodations into budget, mid-range, and luxury tiers."
    )
    add_custom_paragraph(doc, p_desc_fig2)
    
    # Figure 3: Singapore City
    add_figure(doc, "assets/singapore_city.png", "Figure 3: City Explorer - Singapore Page", width_in_inches=5.0)
    p_desc_fig3 = (
        "Figure 3 shows the City Explorer page for 'Singapore'. This view highlights local transit guidelines, "
        "airport details for Changi Airport, and city safety recommendations. It provides international travelers "
        "with essential information on navigating the city-state's public transport systems."
    )
    add_custom_paragraph(doc, p_desc_fig3)
    
    # Figure 4: Tokyo City
    add_figure(doc, "assets/tokyo_city.png", "Figure 4: City Explorer - Tokyo Page", width_in_inches=5.0)
    p_desc_fig4 = (
        "Figure 4 showcases the City Explorer page for 'Tokyo'. It details traditional and modern attractions "
        "(such as Senso-ji Temple and Shibuya Crossing), local dining options, and popular shopping areas "
        "(like Ginza and Akihabara). The layout uses structured cards and custom CSS styling."
    )
    add_custom_paragraph(doc, p_desc_fig4)
    
    doc.add_page_break()
    
    # ==========================================
    # 14. TESTING & QUALITY ASSURANCE
    # ==========================================
    add_custom_heading(doc, "9. Testing & Quality Assurance", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "9.1 BDD-style Pytest Suite", level=2, space_before=12, space_after=6)
    p_test_desc = (
        "The project implements a BDD-style (Behavior Driven Development) unit test suite using Pytest. "
        "The tests are organized into modules verifying specific components of the system. Mocking is utilized "
        "for database connections and API calls to ensure independent, fast test execution. The table below "
        "lists the test suites and the behaviors they verify:"
    )
    add_custom_paragraph(doc, p_test_desc)
    
    table_tests = create_styled_table(doc, 10, 3)
    style_table_header(table_tests.rows[0], ["Test Suite File", "Scope of Verification", "Example Test Cases"])
    
    test_data = [
        ("test_database_spec.py", "Database connection, initialization, and queries", "Verifies database seeding, country/city lookups, and blank search query fallbacks."),
        ("test_auth_spec.py", "User registration, login, and JWT cryptography", "Verifies password hashing, JWT encoding/decoding, and token expiration validation."),
        ("test_chatbot_spec.py", "Chatbot query parsing and context resolution", "Verifies keyword matching, active context fallback, and response formatting."),
        ("test_history_spec.py", "Activity logging and travel history management", "Verifies history item logging, favorite toggling, and history clearing."),
        ("test_i18n_spec.py", "UI translation and internationalization", "Verifies language mapping, key translations, and multi-lingual UI rendering."),
        ("test_styles_spec.py", "CSS injection and HTML rendering utilities", "Verifies custom style injection and markup safety checks."),
        ("test_weather_spec.py", "Weather history lookups and recommendations", "Verifies monthly weather retrieval and packing recommendations."),
        ("test_agent_spec.py", "AI Agent tool calling and Runner operations", "Verifies ADK agent initialization, tool execution, and query coordination."),
        ("test_app_spec.py", "Main router configuration and page loading", "Verifies page routing, sidebar rendering, and session state initialization."),
    ]
    
    for idx, row in enumerate(test_data):
        style_table_row(table_tests.rows[idx + 1], row, is_even=(idx % 2 == 1))
        
    add_custom_heading(doc, "9.2 Quality Assurance and Pre-commit Hooks", level=2, space_before=12, space_after=6)
    p_qa = (
        "To maintain high code quality, security, and formatting consistency, the project enforces strict "
        "pre-commit hooks and CI/CD pipeline checks. These include:"
    )
    add_custom_paragraph(doc, p_qa)
    
    add_bullet_point(doc, "- Code Formatting: ", "Ruff-format is used to maintain a consistent code style across all python files.")
    add_bullet_point(doc, "- Static Analysis & Linting: ", "Ruff, Pylint, Flake8, and Mypy (type-checking) are executed to identify syntax issues, code smells, and type mismatches.")
    add_bullet_point(doc, "- Security Auditing: ", "Bandit scans the codebase for common security flaws (e.g., hardcoded secrets or unsafe imports), and pip-audit checks third-party packages for known vulnerabilities.")
    add_bullet_point(doc, "- Secret Scanning: ", "Gitleaks is integrated to prevent accidental commits of API keys, passwords, or tokens.")
    add_bullet_point(doc, "- CI/CD Pipeline: ", "A GitLab CI/CD pipeline (.gitlab-ci.yml) is configured to run all lints, security scans, and Pytest suites automatically on every commit.")
    
    doc.add_page_break()
    
    # ==========================================
    # 15. ADVANTAGES, LIMITATIONS, & FUTURE SCOPE
    # ==========================================
    add_custom_heading(doc, "10. Advantages, Limitations, & Future Scope", level=1, space_before=24, space_after=12)
    
    add_custom_heading(doc, "10.1 Advantages", level=2, space_before=12, space_after=6)
    add_bullet_point(doc, "- Unified Platform: ", "Consolidates visa checklists, local laws, cultural etiquette, transit guides, hotel listings, and dining options into a single portal.")
    add_bullet_point(doc, "- Dynamic Planning: ", "Generates customized itineraries with interactive cost breakdowns, eliminating static planning limitations.")
    add_bullet_point(doc, "- Context-Aware AI Chatbot: ", "Maintains active country/city selections as conversational context, reducing the need for repetitive user inputs.")
    add_bullet_point(doc, "- Multi-lingual UI: ", "Complete localization in English, Hindi, and Telugu ensures accessibility for a wider range of users.")
    add_bullet_point(doc, "- Secure & Local-First: ", "Uses an embedded SQLite database and local JWT session management, ensuring high performance and privacy.")
    
    add_custom_heading(doc, "10.2 Limitations", level=2, space_before=12, space_after=6)
    add_bullet_point(doc, "- Limited Destination Data: ", "The database is currently pre-seeded with only three countries (India, Japan, Singapore). Expanding the data requires manual seeding or external API integrations.")
    add_bullet_point(doc, "- API Dependency for Advanced AI: ", "The autonomous agent mode requires an active internet connection and a Google Gemini API key; otherwise, the system falls back to a simpler keyword-based matching engine.")
    add_bullet_point(doc, "- Informational Only: ", "The application provides hotel, dining, and transit recommendations but does not support direct bookings or live reservations.")
    
    add_custom_heading(doc, "10.3 Future Scope", level=2, space_before=12, space_after=6)
    add_bullet_point(doc, "- Geographical Expansion: ", "Integrating external travel APIs (e.g., TripAdvisor or Google Places API) to dynamically fetch information for any country or city worldwide.")
    add_bullet_point(doc, "- Live Booking Integration: ", "Partnering with flight and hotel booking APIs (e.g., Amadeus or Booking.com) to allow users to book flights, hotels, and activities directly from the app.")
    add_bullet_point(doc, "- Real-Time Weather Forecasting: ", "Replacing historical weather averages with live weather forecasting APIs (e.g., OpenWeatherMap) to provide real-time packing suggestions.")
    add_bullet_point(doc, "- Voice-Enabled AI Assistant: ", "Integrating speech-to-text and text-to-speech capabilities into the chatbot to assist travelers on the go.")
    add_bullet_point(doc, "- Cross-Platform Mobile Application: ", "Developing a mobile version of TravelMate AI using React Native or Flutter to provide offline access to saved itineraries and emergency guides.")
    
    doc.add_page_break()
    
    # ==========================================
    # 16. CONCLUSION
    # ==========================================
    add_custom_heading(doc, "11. Conclusion", level=1, space_before=24, space_after=12)
    
    p_conclusion = (
        "TravelMate AI successfully addresses the challenges of fragmented information and lack of cultural "
        "awareness in travel planning. By consolidating visa requirements, local regulations, cultural etiquette, "
        "and city recommendations into a single multi-lingual platform, the application provides tourists with "
        "a reliable, structured, and easy-to-use companion. The integration of a dynamic itinerary planner "
        "with interactive budget charts allows for customized trip scheduling. Furthermore, the context-aware "
        "AI chatbot, powered by the Google Agent Development Kit and Gemini, provides intelligent conversational "
        "assistance, while the local rule-based fallback ensures uninterrupted service. Enforced by BDD-style "
        "testing and strict pre-commit quality standards, TravelMate AI is a secure, robust, and scalable "
        "solution that significantly enhances the travel experience, ensuring safety, cultural compliance, and "
        "planning efficiency."
    )
    add_custom_paragraph(doc, p_conclusion)
    
    doc.add_page_break()
    
    # ==========================================
    # 17. REFERENCES
    # ==========================================
    add_custom_heading(doc, "References", level=1, space_before=24, space_after=12)
    
    refs = [
        "[1] Streamlit Documentation, \"Multi-page Apps and Navigation Configuration,\" [Online]. Available: https://docs.streamlit.io.",
        "[2] SQLite Consortium, \"SQLite Database Engine and SQL Syntax Reference,\" [Online]. Available: https://sqlite.org.",
        "[3] Google AI, \"Agent Development Kit (ADK) Developer Guide,\" [Online]. Available: https://github.com/google/adk.",
        "[4] Plotly Technologies, \"Plotly Open Source Graphing Library for Python,\" [Online]. Available: https://plotly.com/python.",
        "[5] Pytest Developer Team, \"Pytest: Helps you write better programs,\" [Online]. Available: https://docs.pytest.org.",
        "[6] Internet Engineering Task Force (IETF), \"RFC 7519: JSON Web Token (JWT) Specification,\" [Online]. Available: https://tools.ietf.org/html/rfc7519.",
        "[7] Python Software Foundation, \"The Python Standard Library - hashlib and hmac modules,\" [Online]. Available: https://docs.python.org/3/library.",
    ]
    
    for r in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.line_spacing = 1.15
        run_ref = p_ref.add_run(r)
        run_ref.font.name = 'Calibri'
        run_ref.font.size = Pt(10)
        run_ref.font.color.rgb = RGBColor(51, 51, 51)
        
    # Save the document
    doc.save(OUTPUT_PATH)
    print(f"Project report successfully generated and saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
